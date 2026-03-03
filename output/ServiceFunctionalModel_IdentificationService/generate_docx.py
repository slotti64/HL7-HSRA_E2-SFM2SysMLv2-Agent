"""Generate a readable illustrative .docx document of the HL7 IS SysML v2 model."""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for lvl in range(1, 5):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# Code style
code_style = doc.styles.add_style('SysMLCode', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(1)

# Placeholder style
ph_style = doc.styles.add_style('DiagramPlaceholder', WD_STYLE_TYPE.PARAGRAPH)
ph_style.font.name = 'Calibri'
ph_style.font.size = Pt(11)
ph_style.font.italic = True
ph_style.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
ph_style.paragraph_format.space_before = Pt(12)
ph_style.paragraph_format.space_after = Pt(12)
ph_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_placeholder(text):
    p = doc.add_paragraph(style='DiagramPlaceholder')
    p.add_run('[DIAGRAM PLACEHOLDER]').bold = True
    doc.add_paragraph(text, style='DiagramPlaceholder')

def add_code(text, max_lines=30):
    lines = text.strip().split('\n')
    if len(lines) > max_lines:
        lines = lines[:max_lines] + [f'    /* ... ({len(lines)-max_lines} more lines) */']
    for line in lines:
        doc.add_paragraph(line, style='SysMLCode')

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('HL7 Identification Service', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_heading('SysML v2 Model Documentation', level=1)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Computation Independent Model (CIM) and Platform Independent Model (PIM)\n\n').bold = True
meta.add_run('Source: HL7 V3 Identification Service (IS) Service Functional Model, Release 1\n')
meta.add_run('Transformation: Automated SFM-to-SysMLv2 Pipeline\n')
meta.add_run('Date: 2026-03-02\n')
meta.add_run('Package: CIM_IS + PIM_IS')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Model Overview',
    '3. CIM: Business Domain Ontology',
    '4. CIM: Stakeholder Model',
    '5. CIM: Business Capabilities (Use Cases)',
    '6. CIM: Business Rules',
    '7. CIM: Requirements',
    '   7.1 Functional Requirements',
    '   7.2 Quality Requirements',
    '   7.3 Compliance Requirements',
    '8. CIM: Traceability',
    '9. PIM: Data Model',
    '10. PIM: Service Contracts',
    '11. PIM: Operations',
    '12. PIM: Behavioral Flows',
    '13. PIM: Composition & State Model',
    '14. PIM: Traceability',
    '15. Transformation Log',
    'Appendix A: Model Statistics',
    'Appendix B: Glossary',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Number 2')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This document presents a comprehensive, human-readable description of the SysML v2 model '
    'derived from the HL7 V3 Identification Service (IS) Service Functional Model (SFM), Release 1. '
    'The model was produced by an automated multi-agent transformation pipeline that analyzes the '
    'SFM specification and generates formal SysML v2 textual notation at two abstraction levels:'
)
doc.add_paragraph('Computation Independent Model (CIM) - business-level concepts, stakeholders, capabilities, rules, and requirements.', style='List Bullet')
doc.add_paragraph('Platform Independent Model (PIM) - data types, service contracts, operations, behavioral workflows, and system composition.', style='List Bullet')
doc.add_paragraph(
    'The Identification Service provides capabilities to manage and retrieve identifying information '
    'for various kinds of entities (people, organizations, devices, etc.) within and across organizations. '
    'It supports identity registration, creation, update, state management, merge/unmerge, link/unlink, '
    'removal, query, and notification operations.'
)

doc.add_heading('1.1 Document Conventions', level=2)
doc.add_paragraph('SysML v2 code excerpts are shown in Consolas monospace font with indentation.', style='List Bullet')
doc.add_paragraph('Diagram placeholders indicate where MagicDraw diagrams should be inserted after model import.', style='List Bullet')
doc.add_paragraph('Statement references [ST-nnn] trace back to the original SFM specification statements.', style='List Bullet')
doc.add_paragraph('Requirement identifiers follow the pattern FR-nnn (Functional), QR-nnn (Quality), CR-nnn (Compliance).', style='List Bullet')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. MODEL OVERVIEW
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('2. Model Overview', level=1)
doc.add_paragraph(
    'The complete model is organized into two top-level SysML v2 packages:'
)
add_table(
    ['Package', 'Level', 'Contents'],
    [
        ['CIM_IS', 'CIM', 'BusinessDomain, StakeholderModel, BusinessCapabilities, BusinessRules, CIM_Requirements, CIM_Traceability'],
        ['PIM_IS', 'PIM', 'DataModel, ServiceContracts, Operations, BehavioralFlows, Composition, PIM_Traceability'],
    ]
)

doc.add_heading('2.1 Package Hierarchy', level=2)
add_code("""CIM_IS/
  BusinessDomain          -- 13 item defs, 14 attribute defs, 5 enums, 12 connection defs, domain model
  StakeholderModel        -- 10 stakeholder part defs (6 primary + 4 secondary actors)
  BusinessCapabilities    -- 29 use cases (15 normative + 14 non-normative scenarios)
  BusinessRules           -- 42 constraint defs (identity, merge, link, query, notification rules)
  CIM_Requirements/
    FunctionalRequirements   -- 41 requirements (FR-001 to FR-041)
    QualityRequirements      -- 11 requirements (QR-001 to QR-011)
    ComplianceRequirements   -- 10 requirements (CR-001 to CR-010)
  CIM_Traceability        -- 77 derivation connections + diagnostic flags

PIM_IS/
  DataModel               -- 45 item defs + 7 enum defs (domain types + request/response pairs)
  ServiceContracts        -- 2 port defs + 2 interface defs + 32 flows
  Operations              -- 15 action defs (1:1 with normative use cases)
  BehavioralFlows         -- 6 workflow action defs (multi-step scenarios)
  Composition             -- 5 part defs + 1 state def + 6 interface connections
  PIM_Traceability        -- 71 dependency statements""")

add_placeholder(
    'Insert: Model Package Hierarchy Diagram\n'
    '(MagicDraw: Package Diagram showing CIM_IS and PIM_IS with sub-packages)'
)

doc.add_heading('2.2 Model Statistics Summary', level=2)
add_table(
    ['Metric', 'Count'],
    [
        ['Total CIM elements', '177'],
        ['Total PIM elements', '170'],
        ['Total requirements', '62 (41 FR + 11 QR + 10 CR)'],
        ['Total traceability links', '225'],
        ['Normative use cases', '15'],
        ['Non-normative business scenarios', '14'],
        ['PIM operations', '15'],
        ['Behavioral workflow flows', '6'],
        ['Business rules', '42 (6 formalized, 36 documentary)'],
        ['Service interfaces', '2 (Management + Query)'],
    ]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. CIM: BUSINESS DOMAIN ONTOLOGY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('3. CIM: Business Domain Ontology', level=1)
doc.add_paragraph(
    'The BusinessDomain package defines the core domain ontology of the Identification Service. '
    'It contains five sections: enumerations, attribute definitions (value types), item definitions '
    '(domain entities), connection definitions (named associations), and a conceptual domain model '
    'with item usages and connection usages. All elements derive from the IS SFM meta-model '
    '(Section 4.2) and glossary (Section 5).'
)

doc.add_heading('3.1 Enumerations', level=2)
doc.add_paragraph(
    'Five enumeration types define the coded value sets used throughout the domain model:'
)
add_table(
    ['Enumeration', 'Values', 'Description', 'Source'],
    [
        ['IdentityStatus', 'ACTIVE, INACTIVE', 'Lifecycle states for identity instances and metadata', 'ST-105, ST-106'],
        ['LinkTypeCategory', 'PEER_LINK, MERGE_DEPRECATION', 'Whether linked identities are peers or one is deprecated', 'ST-220, ST-079'],
        ['LinkMethodCategory', 'MANUAL, AUTOMATED', 'How a link between identities was established', 'ST-221, ST-081'],
        ['UpdateMode', 'OVERWRITE, VALUED, UNSPECIFIED', 'Mode for identity property update operations', 'ST-096'],
        ['MatchQualityTier', 'DEFINITE_MATCH, PRESUMPTIVE_MATCH, NO_MATCH', 'Confidence tier of match results', 'ST-172, ST-205'],
    ]
)

doc.add_heading('3.2 Attribute Definitions (Value Types)', level=2)
doc.add_paragraph(
    'Fourteen abstract attribute definitions represent the value types used as properties '
    'of domain entities. These are intentionally abstract at CIM level and are resolved to '
    'typed String/Boolean/Integer attributes at PIM level.'
)
add_table(
    ['Attribute Def', 'Description', 'Source'],
    [
        ['DomainDesignation', 'Human-readable name for domain concepts', 'ST-069, ST-072, ST-070'],
        ['DomainDescription', 'Textual description of domain concepts', 'ST-069, ST-072, ST-070'],
        ['VersionDesignation', 'Version identifier for domain concepts', 'ST-177'],
        ['LanguageDesignation', 'Human language identifier', 'ST-069'],
        ['IdentityDesignation', 'Unique identity instance identifier within scope', 'ST-091'],
        ['SemanticSignifierContent', 'Populated identifying information structure', 'ST-003, ST-071'],
        ['SchemaSpecification', 'Schema definition for EntityType information structure', 'ST-070, ST-003'],
        ['ValidationRuleSpecification', 'Declarative validation rules for schema', 'ST-070, ST-074'],
        ['ProvenanceDescription', 'Administrative info for link/merge actions', 'ST-080'],
        ['ReasonDescription', 'Reason for management actions', 'ST-077, ST-101'],
        ['MatchConfidence', 'Quality of a match result', 'ST-172, ST-057'],
        ['SubscriberDesignation', 'Notification subscriber identification', 'ST-155, ST-156'],
        ['NotificationQualifier', 'Event type for notifications', 'ST-156, ST-159'],
    ]
)

doc.add_heading('3.3 Domain Entities (Item Definitions)', level=2)
doc.add_paragraph(
    'Thirteen item definitions represent the core domain entities of the Identification Service. '
    'Relationships between entities are modeled as explicit connection definitions (Section 3.4) '
    'rather than embedded references, enabling diagrammable associations in MagicDraw.'
)

entities = [
    ('PolicyDomain', 'An identity space or sphere of use. May be organizational, geographical, or jurisdictional. Defines the scope within which identifiers must be unique.', 'ST-030, ST-069', 'domainDesignation, description, version, status, forCrossReference[0..1], languagesSupported[0..*], defaultLanguage'),
    ('EntityConcept', 'A semantic concept or category (e.g., Person, Device). Enables grouping of different EntityType representations of the same underlying concept.', 'ST-002, ST-072', 'conceptDesignation, description, version'),
    ('EntityType', 'A specific entity type at the level of a specific information model (semantic signifier). Classifies exactly one EntityConcept.', 'ST-071, ST-070', 'typeDesignation, description, status, version, schemaDefinition, validationRuleSet'),
    ('EntityTypeAssignment', 'Bridges an EntityType and a PolicyDomain. Enables domain-specific customization of validation and schema constraints.', 'ST-074', 'status, version, constrainedSchemaDefinition[0..1], constrainedValidationRuleSet[0..1]'),
    ('IdentityInstance', 'The actual identity entry in the service. Holds populated identifying information (semantic signifier values) for a real-world entity.', 'ST-076, ST-091', 'instanceDesignation, version, status, populatedSignifierValues'),
    ('IdentityLink', 'A directed relationship between two identity instances for linking (peers) or merging (one deprecated). Links are transitive.', 'ST-077, ST-079', 'linkType, linkMethod[0..1], reason[0..1], provenance[0..1]'),
    ('RealWorldEntity', 'An actual thing (person, device, organization) that the service manages.', 'ST-002, ST-005', 'entityNature'),
    ('UpdateQualifier', 'Specifies the mode for identity property updates.', 'ST-096', 'updateMode, updateSchemaReference[0..1]'),
    ('SearchQualifier', 'Optional parameters directing search behavior.', 'ST-146, ST-147', 'requestedConfidence[0..1], maximumResultSetSize[0..1], errorIfSizeExceeded[0..1], matchingAlgorithmDesignation[0..1]'),
    ('NotificationSubscription', 'A standing request for identity update notifications.', 'ST-155, ST-159', 'subscriber, notificationQualifier[0..1], subscriptionStatus'),
    ('ServiceMetadata', 'Metadata delineating the scope of a service instance.', 'ST-211', '(attributes defined via connections)'),
    ('HospitalInformationSystem', 'A clinical system acting as a service consumer.', 'ST-045, ST-206', '(no domain attributes)'),
    ('MasterPatientIndex', 'An application providing identity cross-referencing capabilities.', 'ST-207', '(no domain attributes)'),
]
add_table(
    ['Entity', 'Description', 'Source', 'Key Attributes'],
    [(e[0], e[1][:80]+'...' if len(e[1])>80 else e[1], e[2], e[3]) for e in entities]
)

add_placeholder(
    'Insert: Domain Entity Class Diagram\n'
    '(MagicDraw: Block Definition Diagram showing all 13 item defs with attributes)'
)

doc.add_heading('3.4 Domain Associations (Connection Definitions)', level=2)
doc.add_paragraph(
    'Twelve connection definitions model the named relationships from the SFM meta-model (Section 4.2). '
    'Each connection has typed end features with cardinalities reflecting the specification constraints.'
)

doc.add_heading('3.4.1 Core Meta-Model Associations', level=3)
add_table(
    ['Connection', 'From', 'To', 'Description', 'Source'],
    [
        ['Classifies', 'EntityType [0..*]', 'EntityConcept [1]', 'An EntityType classifies exactly one EntityConcept', 'ST-073'],
        ['DefinesAssignment', 'EntityType [1]', 'EntityTypeAssignment [0..*]', 'An EntityType has zero or more domain-specific assignments', 'ST-074'],
        ['IsAppliedTo', 'EntityTypeAssignment [0..*]', 'PolicyDomain [1]', 'An assignment applies within exactly one PolicyDomain', 'ST-074'],
        ['ProvidesRuleFor', 'IdentityInstance [0..*]', 'EntityTypeAssignment [1]', 'An instance is governed by exactly one assignment', 'ST-076'],
        ['LinkSource', 'IdentityLink [0..*]', 'IdentityInstance [1]', 'Source identity in a link (deprecated in merge)', 'ST-077'],
        ['LinkTarget', 'IdentityLink [0..*]', 'IdentityInstance [1]', 'Target identity in a link (surviving in merge)', 'ST-077'],
    ]
)

doc.add_heading('3.4.2 Subscription Filter & Service Metadata Associations', level=3)
add_table(
    ['Connection', 'From', 'To', 'Description'],
    [
        ['SubscriptionDomainFilter', 'NotificationSubscription [0..*]', 'PolicyDomain [0..1]', 'Optional PolicyDomain filter for subscription'],
        ['SubscriptionTypeFilter', 'NotificationSubscription [0..*]', 'EntityType [0..1]', 'Optional EntityType filter for subscription'],
        ['SubscriptionInstanceFilter', 'NotificationSubscription [0..*]', 'IdentityInstance [0..1]', 'Optional specific instance filter'],
        ['SupportsDomain', 'ServiceMetadata [0..*]', 'PolicyDomain [1..*]', 'PolicyDomains managed by service instance'],
        ['SupportsEntityType', 'ServiceMetadata [0..*]', 'EntityType [1..*]', 'EntityTypes recognized by service instance'],
        ['SupportsConcept', 'ServiceMetadata [0..*]', 'EntityConcept [1..*]', 'EntityConcepts recognized by service instance'],
    ]
)

add_placeholder(
    'Insert: Domain Ontology Interconnection Diagram\n'
    '(MagicDraw: Internal Block Diagram of domainModel showing all 13 item usages\n'
    'and 12 connection usages with association labels and cardinalities)'
)

doc.add_heading('3.5 Domain Model (Conceptual Instance)', level=2)
doc.add_paragraph(
    'The domain model section instantiates all 13 entity types as item usages and all 12 '
    'associations as connection usages within a single part definition. This creates a '
    'diagrammable conceptual model that can be rendered as an interconnection diagram in MagicDraw.'
)
add_code("""part domainModel {
    /* Item Usages */
    item policyDomains : PolicyDomain [1..*];
    item entityConcepts : EntityConcept [1..*];
    item entityTypes : EntityType [1..*];
    item entityTypeAssignments : EntityTypeAssignment [0..*];
    item identityInstances : IdentityInstance [0..*];
    item identityLinks : IdentityLink [0..*];
    /* ... 7 more item usages ... */

    /* Connection Usages */
    connection classifies : Classifies connect entityTypes to entityConcepts;
    connection definesAssignment : DefinesAssignment connect entityTypes to entityTypeAssignments;
    connection isAppliedTo : IsAppliedTo connect entityTypeAssignments to policyDomains;
    connection providesRuleFor : ProvidesRuleFor connect identityInstances to entityTypeAssignments;
    /* ... 8 more connection usages ... */
}""")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. CIM: STAKEHOLDER MODEL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('4. CIM: Stakeholder Model', level=1)
doc.add_paragraph(
    'The StakeholderModel package defines all business stakeholders who interact with the '
    'Identification Service. Stakeholders are derived from SFM Section 3 (Business Scenarios) '
    'primary actors and from roles implied in Section 4 (Detailed Functional Model).'
)

doc.add_heading('4.1 Primary Actors', level=2)
add_table(
    ['Stakeholder', 'Description', 'Key Concern', 'Source'],
    [
        ['RegistrationClerk', 'Hospital clerk for admissions', 'Efficient search, create, update, link, merge', 'ST-212'],
        ['ClinicalUser', 'Healthcare professional (nurse, physician)', 'Quick/accurate identity lookup', 'ST-213'],
        ['LaboratoryClerk', 'Lab staff processing orders', 'Reliable identifier resolution including merged IDs', 'ST-046'],
        ['IdentityResolutionClerk', 'Specialized staff for merge/unmerge/link/unlink', 'Authority to resolve duplicate identities', 'ST-214'],
        ['ServiceAdministrator', 'System administrator', 'Manage configuration; access all records incl. inactive', 'ST-215'],
        ['ExternalSystem', 'External application (referral, lab, radiology)', 'Programmatic identity resolution', 'ST-216'],
    ]
)

doc.add_heading('4.2 Secondary Actors', level=2)
add_table(
    ['Stakeholder', 'Description', 'Source'],
    [
        ['PolicyDomainOwner', 'Owner who defines identity policies for a domain', 'ST-030, ST-180'],
        ['Patient', 'Individual who is the subject of identity management', 'ST-048, ST-002'],
        ['ReferralSystem', 'System receiving referral info from external institutions', 'ST-187'],
        ['RegionalHealthNetwork', 'Multi-org network operating cross-domain IS (XIS)', 'ST-195'],
    ]
)

add_placeholder(
    'Insert: Stakeholder Diagram\n'
    '(MagicDraw: Block Definition Diagram showing BusinessStakeholder hierarchy\n'
    'with 6 primary and 4 secondary actor part defs)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. CIM: BUSINESS CAPABILITIES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('5. CIM: Business Capabilities (Use Cases)', level=1)
doc.add_paragraph(
    'The BusinessCapabilities package models all service capabilities as use case usages, organized '
    'into a top-level service capability with two normative capability groups (Identity Management '
    'and Query) plus non-normative business scenarios from SFM Section 3.'
)

doc.add_heading('5.1 Identity Management Capabilities (Section 4.3)', level=2)
doc.add_paragraph('Nine normative operations for creating, modifying, and removing identity data:')
add_table(
    ['Use Case', 'Description', 'Source'],
    [
        ['registerIdentity', 'Create identity with consumer-supplied unique identifier', 'ST-082'],
        ['createIdentity', 'Create identity with service-generated identifier', 'ST-088'],
        ['updateIdentityProperties', 'Update identifying property values with update qualifier modes', 'ST-092'],
        ['updateIdentityState', 'Change processing state (ACTIVE/INACTIVE)', 'ST-097'],
        ['mergeIdentities', 'Merge two identities within single PolicyDomain', 'ST-107'],
        ['unmergeIdentities', 'Reverse a previous merge', 'ST-115'],
        ['linkIdentities', 'Establish link between identities (within or across domains)', 'ST-121'],
        ['unlinkIdentity', 'Remove a previously established link', 'ST-129'],
        ['removeIdentityInstance', 'Permanently remove an identity instance', 'ST-133'],
    ]
)

doc.add_heading('5.2 Query Capabilities (Section 4.4)', level=2)
doc.add_paragraph('Six normative operations for retrieving identity information and managing notifications:')
add_table(
    ['Use Case', 'Description', 'Source'],
    [
        ['getAllInformationForIdentity', 'Retrieve all info for a known identity; redirect if merged', 'ST-140'],
        ['findIdentitiesByProperty', 'Search by partial property values with match quality', 'ST-146'],
        ['listLinkedIdentities', 'List identities linked to a specified identity', 'ST-150'],
        ['requestIdentityUpdateNotifications', 'Subscribe to identity change notifications', 'ST-155'],
        ['updateNotificationSubscription', 'Modify or cancel a notification subscription', 'ST-160'],
        ['notifyIdentityUpdates', 'Publish notifications to matching subscribers', 'ST-164'],
    ]
)

add_placeholder(
    'Insert: Use Case Diagram (Normative Operations)\n'
    '(MagicDraw: Use Case Diagram showing identificationService with 15 included use cases,\n'
    'grouped into IdentityManagement and Query capability groups, with actor associations)'
)

doc.add_heading('5.3 Business Scenarios (Non-Normative)', level=2)
doc.add_paragraph(
    'Fourteen non-normative business scenarios from SFM Section 3 illustrate how the normative '
    'operations are composed in realistic clinical workflows. These scenarios are not requirements '
    'for conformance (per ST-202) but serve as explanatory examples.'
)
doc.add_heading('5.3.1 Single-Domain Scenarios', level=3)
add_table(
    ['Scenario', 'Description', 'Composes Operations', 'Source'],
    [
        ['createNewPatient', 'Carol searches, not found, registers new identity', 'findByProperty, registerIdentity', 'ST-048'],
        ['linkOrMergeEntities', 'Nancy finds duplicates, links or merges per policy', 'findByProperty, linkIdentities/mergeIdentities', 'ST-052'],
        ['updateDemographics', 'Carol searches, retrieves, updates address', 'findByProperty, getAllInfo, updateProperties', 'ST-056'],
        ['inactivateEntity', 'Discharge nurse inactivates patient', 'updateIdentityState', 'ST-059'],
        ['activateEntity', 'Re-admit and reactivate patient', 'updateIdentityState', 'ST-062'],
        ['unlinkEntity', 'Discover and remove erroneous link', 'findByProperty, listLinked, unlinkIdentity', 'ST-064'],
        ['lookupSingleEntry', 'Search returns single definite match', 'findByProperty', 'ST-181'],
        ['lookupMultipleEntries', 'Search returns multiple candidates', 'findByProperty', 'ST-183'],
        ['mergedEntriesFound', 'Retrieve merged identity, get redirect', 'getAllInfo', 'ST-185'],
        ['unattendedEncounter', 'Referral system auto-resolves identity', 'findByProperty, createIdentity', 'ST-187'],
        ['removeEntity', 'Administrator removes identity', 'removeIdentityInstance', 'ST-192'],
    ]
)
doc.add_heading('5.3.2 Multi-Domain Scenarios', level=3)
add_table(
    ['Scenario', 'Description', 'Source'],
    [
        ['lookupAcrossRegionalNetwork', 'Search across all member organizations via XIS', 'ST-196'],
        ['lookupAtSpecificOrganization', 'Directed search at specific member organization', 'ST-198'],
        ['linkEntitiesAcrossRegions', 'Register local identity and link to regional master', 'ST-200'],
    ]
)

add_placeholder(
    'Insert: Business Scenario Use Case Diagram\n'
    '(MagicDraw: Use Case Diagram for singleDomainScenarios and multiDomainScenarios\n'
    'showing actors Carol Clerk, Nancy Nightingale, Eric Entry, Bill Beaker, etc.)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. CIM: BUSINESS RULES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('6. CIM: Business Rules', level=1)
doc.add_paragraph(
    'The BusinessRules package contains 42 constraint definitions derived from the IS SFM Section 4 '
    'normative statements and related policy declarations. Rules are organized into functional categories.'
)

rule_categories = [
    ('Identity Uniqueness & Scope', [
        ('IdentityUniquenessWithinDomain', 'Identity identifier must be unique within PolicyDomain + EntityType', 'Formalized'),
        ('SingleDomainInteraction', 'Individual interactions operate within one PolicyDomain', 'Documentary'),
        ('PolicyDomainOwnershipRequired', 'Each PolicyDomain must have an owner', 'Documentary'),
    ]),
    ('Entity Type & Semantic Signifier', [
        ('UniformInterfaceBehavior', 'Interface behavior is invariant across EntityTypes', 'Documentary'),
        ('EntityTypeSpecializationRequired', 'Semantic signifier must be specialized per entity type', 'Documentary'),
        ('EntityTypeClassifiesOneConcept', 'Each EntityType classifies exactly one EntityConcept', 'Formalized'),
        ('AssignmentConstrainsEntityType', 'Assignment may only further constrain EntityType rules', 'Documentary'),
    ]),
    ('Merge', [
        ('MergeWithinSingleDomainOnly', 'Merge restricted to single PolicyDomain', 'Formalized'),
        ('MergeSameEntityConceptRequired', 'Merge restricted to same EntityConcept', 'Formalized'),
        ('MergeAttributeResolution', 'Empty target attributes filled from source during merge', 'Documentary'),
        ('MergeIndicationRequired', 'Merge must be indicated (link or correlation set)', 'Documentary'),
    ]),
    ('Link', [
        ('LinkPreconditions', 'Both identities must exist; target must be active', 'Formalized'),
        ('LinkSemanticsAndTransitivity', 'Links are transitive by definition', 'Documentary'),
        ('CrossEntityTypeLinkingAllowed', 'Different EntityTypes may be linked', 'Documentary'),
        ('CrossDomainLinkTopology', 'Peer-to-peer or master-domain per policy', 'Documentary'),
    ]),
    ('Query & Notification', [
        ('MergedIdentityRedirect', 'Warning + surviving ID on merged lookup', 'Documentary'),
        ('GetAllInfoInvariant', 'Query is read-only', 'Documentary'),
        ('TwoTieredMatchThreshold', 'Definite vs. presumptive match tiers', 'Documentary'),
        ('NotificationSubscriptionPattern', 'Subscription is a pub/sub pattern', 'Documentary'),
    ]),
]

for cat_name, rules in rule_categories:
    doc.add_heading(f'6.x {cat_name} Rules', level=2)
    add_table(
        ['Rule', 'Description', 'Formalization'],
        rules
    )

doc.add_paragraph(
    'Additionally, 20 rules cover update operations, state changes, unmerge, unlink, remove, '
    'and cross-cutting policy concerns (security delegation, authorization, demographic data guidance, '
    'versioning, profile-based constraining, etc.).'
)

add_placeholder(
    'Insert: Business Rules Dependency Diagram\n'
    '(MagicDraw: Requirement Diagram showing constraint defs grouped by category\n'
    'with derivation links to functional requirements)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. CIM: REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('7. CIM: Requirements', level=1)
doc.add_paragraph(
    'The CIM_Requirements package contains 62 INCOSE-compliant requirements organized into '
    'three sub-packages: Functional (41), Quality (11), and Compliance (10).'
)

doc.add_heading('7.1 Functional Requirements', level=2)
doc.add_paragraph(
    'Forty-one functional requirements define what the system must do. They are grouped by '
    'operational capability area.'
)

fr_groups = [
    ('Identity Registration & Creation', [
        ('FR-001', 'Register identity with supplied identifier'),
        ('FR-002', 'Validate supplied identifier uniqueness'),
        ('FR-003', 'Validate properties against schema'),
        ('FR-004', 'Create identity with generated identifier'),
        ('FR-005', 'Trigger automated implicit linking'),
    ]),
    ('Identity Update', [
        ('FR-006', 'Update identity property values'),
        ('FR-007', 'Apply update qualifier modes (OVERWRITE, VALUED, UNSPECIFIED)'),
        ('FR-008', 'Update identity processing state'),
        ('FR-009', 'Enforce state transition validity'),
        ('FR-010', 'Exclude inactive identities from clinical searches'),
        ('FR-011', 'Allow administrators to find inactive identities'),
    ]),
    ('Merge & Unmerge', [
        ('FR-012', 'Merge two identities within single domain'),
        ('FR-013', 'Restrict merge to same entity concept'),
        ('FR-014', 'Resolve attributes during merge'),
        ('FR-015', 'Record merge indication'),
        ('FR-016', 'Unmerge previously merged identities'),
        ('FR-017', 'Flag unmerged identities for manual review'),
    ]),
    ('Link & Unlink', [
        ('FR-018', 'Link identities within or across domains'),
        ('FR-019', 'Enforce link transitivity'),
        ('FR-020', 'Permit cross-entity-type linking'),
        ('FR-021', 'Unlink previously linked identities'),
    ]),
    ('Remove', [
        ('FR-022', 'Remove identity instance'),
        ('FR-023', 'Prevent removal of linked identities'),
        ('FR-024', 'Discourage identifier reuse after removal'),
    ]),
    ('Query Operations', [
        ('FR-025', 'Retrieve all information for known identity'),
        ('FR-026', 'Redirect to surviving identity on merged lookup'),
        ('FR-027', 'Find identities by property'),
        ('FR-028', 'Return match quality with search results'),
        ('FR-029', 'List linked identities'),
    ]),
    ('Notification', [
        ('FR-030', 'Subscribe to identity update notifications'),
        ('FR-031', 'Update or cancel notification subscription'),
        ('FR-032', 'Publish identity update notifications'),
    ]),
    ('Cross-Cutting', [
        ('FR-033', 'Support multiple entity types'),
        ('FR-034', 'Associate domain-specific identifiers in XIS'),
        ('FR-035', 'Provide signifier discoverability'),
        ('FR-036', 'Enforce uniform interface behavior'),
        ('FR-037', 'Scope interactions to single policy domain'),
        ('FR-038', 'Report operation exceptions'),
        ('FR-039', 'Retain removed records for audit (placeholder)'),
        ('FR-040', 'Manage service metadata (placeholder)'),
        ('FR-041', 'Support profile-based constraining'),
    ]),
]

for group_name, reqs in fr_groups:
    doc.add_heading(f'7.1.x {group_name}', level=3)
    add_table(['ID', 'Requirement'], reqs)

add_placeholder(
    'Insert: Functional Requirements Traceability Matrix\n'
    '(MagicDraw: Requirements Table showing FR-001 to FR-041 with derivation links\n'
    'to use cases and business rules)'
)

doc.add_heading('7.2 Quality Requirements', level=2)
doc.add_paragraph(
    'Eleven quality requirements are derived from gap analysis (the IS SFM contains no explicit '
    'quality-of-service requirements). No quantified thresholds are specified at CIM level.'
)
add_table(
    ['ID', 'Requirement', 'Quality Attribute'],
    [
        ('QR-001', 'Service availability during published service hours', 'Availability'),
        ('QR-002', 'Search response timeliness', 'Performance'),
        ('QR-003', 'Identity retrieval timeliness', 'Performance'),
        ('QR-004', 'Scalability for identity volume', 'Scalability'),
        ('QR-005', 'Concurrent access integrity', 'Reliability'),
        ('QR-006', 'Notification delivery reliability', 'Reliability'),
        ('QR-007', 'Match algorithm effectiveness (two-tier)', 'Accuracy'),
        ('QR-008', 'Version consistency', 'Data Integrity'),
        ('QR-009', 'Language support within policy domains', 'Usability'),
        ('QR-010', 'Identity uniqueness guarantee', 'Data Integrity'),
        ('QR-011', 'Read-only query operations', 'Data Integrity'),
    ]
)

doc.add_heading('7.3 Compliance Requirements', level=2)
add_table(
    ['ID', 'Requirement', 'Regulatory Source'],
    [
        ('CR-001', 'ISO 21090 data type compliance', 'ISO 21090:2011'),
        ('CR-002', 'Authentication precondition on all operations', 'Org security policy (ST-171)'),
        ('CR-003', 'Authorization enforcement for operations', 'Org authorization policy (ST-067)'),
        ('CR-004', 'Restrict identity resolution privileges', 'Org role-based access (ST-067)'),
        ('CR-005', 'Restrict removal privileges', 'Org access control (ST-139)'),
        ('CR-006', 'Audit trail for removed records', 'Org data retention (ST-194)'),
        ('CR-007', 'Policy domain ownership enforcement', 'Org governance (ST-030)'),
        ('CR-008', 'Conformance through profiles', 'HL7 conformance framework'),
        ('CR-009', 'Normative section compliance', 'HL7 IS SFM Section 4'),
        ('CR-010', 'Entity type assignment constraints', 'HL7 IS SFM meta-model (ST-074)'),
    ]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 8. CIM: TRACEABILITY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('8. CIM: Traceability', level=1)
doc.add_paragraph(
    'The CIM_Traceability package contains 77 #derivation connection statements linking '
    'requirements to their originating use cases and business rules. It also includes '
    'diagnostic flags for completeness auditing.'
)
add_table(
    ['Traceability Metric', 'Value'],
    [
        ['Total #derivation connections', '77'],
        ['Use case coverage', '15/15 normative use cases linked (100%)'],
        ['Rule coverage', '22/42 rules linked directly (55%)'],
        ['Orphan requirements', '5 (FR-040, CR-001, CR-007, CR-008, CR-009) - all justified'],
        ['Unjustified capabilities', '0 normative, 14 non-normative (by design)'],
        ['Unrefined rules', '20 (11 exception-list rules + 9 guidance/meta)'],
    ]
)

add_placeholder(
    'Insert: CIM Traceability Matrix Diagram\n'
    '(MagicDraw: Requirement Derivation Diagram showing use cases -> requirements\n'
    'derivation connections with color-coded links by category)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 9. PIM: DATA MODEL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('9. PIM: Data Model', level=1)
doc.add_paragraph(
    'The PIM DataModel package contains platform-independent data type definitions derived from '
    'CIM BusinessDomain entities. It adds technical identifiers (String-typed IDs), typed attributes '
    '(resolving CIM abstract attribute defs to String/Boolean/Integer), and Request/Response pairs '
    'for all 15 operations.'
)

doc.add_heading('9.1 Domain Entity Types', level=2)
doc.add_paragraph(
    'Fourteen domain entity item defs mirror CIM BusinessDomain entities with added technical identifiers:'
)
add_table(
    ['PIM Type', 'CIM Origin', 'Added at PIM', 'Key Attributes'],
    [
        ['PolicyDomain', 'BusinessDomain::PolicyDomain', 'policyDomainId : String', 'domainDesignation, description, version, status, forCrossReference, languagesSupported, defaultLanguage'],
        ['EntityConcept', 'BusinessDomain::EntityConcept', 'entityConceptId : String', 'conceptDesignation, description, version'],
        ['EntityType', 'BusinessDomain::EntityType', 'entityTypeId : String', 'typeDesignation, description, status, version, schemaDefinition, validationRuleSet, ref classifiedConcept'],
        ['IdentityInstance', 'BusinessDomain::IdentityInstance', 'identityInstanceId : String', 'instanceDesignation, version, status, populatedSignifierValues, ref governingAssignment'],
        ['IdentityLink', 'BusinessDomain::IdentityLink', 'identityLinkId : String', 'linkType, linkMethod, reason, provenance, ref sourceInstance, ref targetInstance'],
        ['ServiceFault', '(new at PIM)', 'faultCode : String', 'faultMessage, faultCategory (VALIDATION/AUTHORIZATION/NOT_FOUND/BUSINESS_RULE_VIOLATION/INTERNAL)'],
    ]
)

doc.add_heading('9.2 Request/Response Pairs', level=2)
doc.add_paragraph(
    'Thirty Request/Response item defs (15 pairs) provide typed payloads for each operation:'
)
add_table(
    ['Operation', 'Request Type', 'Response Type', 'Key Request Fields'],
    [
        ['RegisterIdentity', 'RegisterIdentityRequest', 'RegisterIdentityResponse', 'suppliedIdentifier, policyDomainId, entityTypeId, populatedSignifierValues'],
        ['CreateIdentity', 'CreateIdentityRequest', 'CreateIdentityResponse', 'policyDomainId, entityTypeId, populatedSignifierValues'],
        ['UpdateIdentityProperties', 'UpdateIdentityPropertiesRequest', 'UpdateIdentityPropertiesResponse', 'identityInstanceId, updatedSignifierValues, updateMode'],
        ['MergeIdentities', 'MergeIdentitiesRequest', 'MergeIdentitiesResponse', 'sourceIdentityId, targetIdentityId, policyDomainId'],
        ['LinkIdentities', 'LinkIdentitiesRequest', 'LinkIdentitiesResponse', 'sourceIdentityId, targetIdentityId (cross-domain fields)'],
        ['FindIdentitiesByProperty', 'FindIdentitiesByPropertyRequest', 'FindIdentitiesByPropertyResponse', 'searchSignifierValues, requestedConfidence, maximumResultSetSize'],
        ['...', '(9 more pairs)', '...', '...'],
    ]
)

add_placeholder(
    'Insert: PIM Data Model Class Diagram\n'
    '(MagicDraw: Block Definition Diagram showing all PIM item defs organized\n'
    'into Domain Entities, Enumerations, and Request/Response sections)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. PIM: SERVICE CONTRACTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('10. PIM: Service Contracts', level=1)
doc.add_paragraph(
    'The ServiceContracts package defines the logical service boundaries. Per SFM Section 2.3.1, '
    'two interfaces are defined:'
)

doc.add_heading('10.1 Port Definitions', level=2)
add_table(
    ['Port', 'Operations', 'Direction Pattern'],
    [
        ['IdentificationManagementPort', '9 mutation operations + fault', 'out request / in response per operation'],
        ['QueryPort', '6 query/notification operations + fault', 'out request / in response per operation'],
    ]
)

doc.add_heading('10.2 Interface Definitions', level=2)
add_table(
    ['Interface', 'End: requester', 'End: provider', 'Flows'],
    [
        ['IdentificationManagementAPI', '~IdentificationManagementPort', 'IdentificationManagementPort', '19 flows (9 request + 9 response + 1 fault)'],
        ['QueryAPI', '~QueryPort', 'QueryPort', '13 flows (6 request + 6 response + 1 fault)'],
    ]
)
doc.add_paragraph(
    'Each flow is a typed ref flow with explicit "from requester.X to provider.X" direction for '
    'requests and "from provider.Y to requester.Y" for responses, using fully qualified PIM_IS::DataModel types.'
)

add_placeholder(
    'Insert: Service Contract Interface Diagram\n'
    '(MagicDraw: Interface Definition Diagram showing IdentificationManagementAPI and QueryAPI\n'
    'with port types, conjugation (~), and flow definitions)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 11. PIM: OPERATIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('11. PIM: Operations', level=1)
doc.add_paragraph(
    'The Operations package contains 15 action definitions, one for each normative CIM use case. '
    'Each operation has typed in/out parameters, documented preconditions, and documented postconditions.'
)

add_table(
    ['Operation', 'Interface', 'Satisfies', 'Key Preconditions', 'Key Postconditions'],
    [
        ['RegisterIdentity', 'Management', 'FR-001..005', 'ID not exists, schema valid, auth', 'New instance ACTIVE, implicit linking evaluated'],
        ['CreateIdentity', 'Management', 'FR-003..005', 'Domain/type recognized, schema valid', 'New instance with generated ID, ACTIVE'],
        ['UpdateIdentityProperties', 'Management', 'FR-006, FR-007', 'Instance exists, schema valid', 'Properties updated per mode, version incremented'],
        ['UpdateIdentityState', 'Management', 'FR-008..011', 'Valid state transition', 'State changed, other info unchanged'],
        ['MergeIdentities', 'Management', 'FR-012..015', 'Same domain, same concept, valid states', 'Source deprecated, target enriched'],
        ['UnmergeIdentities', 'Management', 'FR-016, FR-017', 'Previously merged', 'Source reinstated, flagged for review'],
        ['LinkIdentities', 'Management', 'FR-018..020', 'Both exist, target active', 'Link created, transitivity applies'],
        ['UnlinkIdentity', 'Management', 'FR-021', 'Currently linked', 'Link removed'],
        ['RemoveIdentityInstance', 'Management', 'FR-022..024', 'Exists, not linked, privileged', 'No longer discoverable, audit retained'],
        ['GetAllInformation', 'Query', 'FR-025, FR-026', 'ID recognized', 'Read-only, redirect if merged'],
        ['FindIdentitiesByProperty', 'Query', 'FR-027, FR-028', 'Valid search params', 'Results with match quality tiers'],
        ['ListLinkedIdentities', 'Query', 'FR-029', 'ID exists', 'All linked IDs returned, transitive'],
        ['RequestIdentityUpdateNotifications', 'Query', 'FR-030', 'Valid criteria', 'Subscription ACTIVE'],
        ['UpdateNotificationSubscription', 'Query', 'FR-031', 'Subscription exists', 'Updated or cancelled'],
        ['NotifyIdentityUpdates', 'Query', 'FR-032', 'Change detected, subscriptions match', 'Notifications delivered'],
    ]
)

add_placeholder(
    'Insert: Operations Activity Diagram\n'
    '(MagicDraw: Activity Diagram showing all 15 action defs with their\n'
    'in/out parameters grouped by Management vs Query interface)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 12. PIM: BEHAVIORAL FLOWS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('12. PIM: Behavioral Flows', level=1)
doc.add_paragraph(
    'Six behavioral flow action defs model multi-step workflow orchestrations derived from '
    'CIM business scenarios that compose multiple normative PIM operations. Only workflows '
    'with genuine multi-operation composition are modeled; single-operation scenarios are not given flows.'
)

flows = [
    ('CreateNewPatientFlow', 'ST-048', 'FindIdentitiesByProperty -> FindIdentitiesByProperty -> RegisterIdentity',
     'Carol searches by name/address (not found), searches with additional info (not found), then registers.'),
    ('LinkOrMergeEntitiesFlow', 'ST-052', 'FindIdentitiesByProperty -> [decide] -> LinkIdentities | MergeIdentities',
     'Nancy finds duplicates, then based on hospital policy either links or merges them.'),
    ('UpdateDemographicsFlow', 'ST-056', 'FindIdentitiesByProperty -> GetAllInformation -> UpdateIdentityProperties',
     'Carol searches, retrieves full record, verifies identity, then updates properties.'),
    ('UnlinkEntityFlow', 'ST-064', 'FindIdentitiesByProperty -> ListLinkedIdentities -> UnlinkIdentity',
     'Carol discovers erroneous link, retrieves alias list, reviews and requests unlink.'),
    ('UnattendedEncounterFlow', 'ST-187', 'FindIdentitiesByProperty -> [decide] -> CreateIdentity | complete | await',
     'Referral system searches IS: definite match (complete), no match (auto-create or await), presumptive (await).'),
    ('CrossRegionLinkFlow', 'ST-200', 'FindIdentitiesByProperty -> RegisterIdentity -> LinkIdentities',
     'Eric searches across regions via XIS, registers local identity, then links local to master.'),
]

for name, src, composition, desc in flows:
    doc.add_heading(f'12.x {name}', level=2)
    doc.add_paragraph(f'Source: {src}')
    doc.add_paragraph(f'Composition: {composition}')
    doc.add_paragraph(desc)

add_placeholder(
    'Insert: Behavioral Flow Activity Diagrams (one per flow)\n'
    '(MagicDraw: Activity Diagrams for each of the 6 flows showing\n'
    'action nodes, decision/merge nodes, and succession edges.\n'
    'Recommended: one diagram per flow for readability.)\n\n'
    'Diagrams needed:\n'
    '  - Fig 12.1: CreateNewPatientFlow (sequential, 3 steps)\n'
    '  - Fig 12.2: LinkOrMergeEntitiesFlow (decision branch)\n'
    '  - Fig 12.3: UpdateDemographicsFlow (sequential, 3 steps)\n'
    '  - Fig 12.4: UnlinkEntityFlow (sequential, 3 steps)\n'
    '  - Fig 12.5: UnattendedEncounterFlow (3-way decision)\n'
    '  - Fig 12.6: CrossRegionLinkFlow (sequential, 3 steps)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 13. PIM: COMPOSITION & STATE MODEL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('13. PIM: Composition & State Model', level=1)

doc.add_heading('13.1 System Composition', level=2)
doc.add_paragraph(
    'The Composition package defines the logical architecture with a single unified provider '
    'exposing both Management and Query endpoints, and three consumer categories derived from '
    'SFM Section 3 business scenarios.'
)

doc.add_heading('13.1.1 Provider', level=3)
doc.add_paragraph(
    'IdentificationServiceProvider exposes two port endpoints:'
)
add_table(
    ['Port', 'Type', 'Operations Exposed'],
    [
        ['managementEndpoint', 'IdentificationManagementPort', '9 mutation operations'],
        ['queryEndpoint', 'QueryPort', '6 query/notification operations'],
    ]
)

doc.add_heading('13.1.2 Consumers', level=3)
add_table(
    ['Consumer', 'Description', 'CIM Stakeholders', 'Ports'],
    [
        ['ClinicalSystemConsumer', 'HIS operated by clerks/nurses/admins', 'RegistrationClerk, ClinicalUser, LaboratoryClerk, ServiceAdministrator', '~ManagementPort + ~QueryPort'],
        ['ReferralSystemConsumer', 'Unattended referral system', 'ExternalSystem, ReferralSystem', '~ManagementPort + ~QueryPort'],
        ['FederatedISConsumer', 'XIS/RHIO intermediary', 'RegionalHealthNetwork', '~ManagementPort + ~QueryPort'],
    ]
)

doc.add_heading('13.1.3 System-Level Connections', level=3)
doc.add_paragraph(
    'The IdentificationServiceSystem top-level part definition connects all consumers to the provider '
    'using typed interface connections (6 total: 2 per consumer).'
)

add_placeholder(
    'Insert: System Composition Internal Block Diagram\n'
    '(MagicDraw: Internal Block Diagram of IdentificationServiceSystem showing\n'
    'provider with managementEndpoint and queryEndpoint ports,\n'
    '3 consumers each with conjugated ports,\n'
    'and 6 interface connections with typed API references)'
)

doc.add_heading('13.2 Identity Lifecycle State Machine', level=2)
doc.add_paragraph(
    'A state definition models the lifecycle of identity instances with three states and five transitions:'
)
add_table(
    ['State', 'Description', 'Entry Condition'],
    [
        ['active', 'Discoverable in clinical searches', 'After registration/creation or reactivation'],
        ['inactive', 'Excluded from clinical searches, visible to admins', 'After inactivation'],
        ['deprecated', 'Merged into surviving identity', 'After merge (source identity)'],
    ]
)
add_table(
    ['Transition', 'From', 'To', 'Trigger'],
    [
        ['creation', 'initial', 'active', 'RegisterIdentity or CreateIdentity'],
        ['inactivation', 'active', 'inactive', 'UpdateIdentityState (requestedStatus=INACTIVE)'],
        ['reactivation', 'inactive', 'active', 'UpdateIdentityState (requestedStatus=ACTIVE)'],
        ['mergeDeprecation', 'active', 'deprecated', 'MergeIdentities (source identity)'],
        ['unmergeReinstatement', 'deprecated', 'active', 'UnmergeIdentities (requires manual review)'],
    ]
)

add_placeholder(
    'Insert: Identity Lifecycle State Machine Diagram\n'
    '(MagicDraw: State Machine Diagram for IdentityLifecycleState showing\n'
    'initial -> active, active <-> inactive, active -> deprecated -> active\n'
    'with transition names and trigger annotations)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 14. PIM: TRACEABILITY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('14. PIM: Traceability', level=1)
doc.add_paragraph(
    'The PIM_Traceability package contains 71 dependency statements mapping every PIM element '
    'back to its CIM antecedent. Traceability covers nine categories:'
)
add_table(
    ['Category', 'Count', 'Direction'],
    [
        ['PIM Operations -> CIM Use Cases', '15', 'realizes'],
        ['PIM Domain Data Types -> CIM Business Domain', '11', 'derives from'],
        ['PIM-Derived Data Types -> CIM Supporting Concepts', '3', 'derives from'],
        ['PIM Enumerations -> CIM Domain Enumerations', '7', 'derives from'],
        ['PIM Request/Response Types -> CIM Use Case I/O', '30', 'derives from'],
        ['PIM Service Contracts -> CIM Capability Groups', '4', 'exposes'],
        ['PIM Behavioral Flows -> CIM Business Scenarios', '6', 'orchestrates'],
        ['PIM Composition Parts -> CIM Stakeholders', '5', 'serves'],
        ['PIM State Model -> CIM Rules + Domain Concepts', '1', 'implements'],
    ]
)
doc.add_paragraph(
    'Completeness self-check: all PIM elements have at least one CIM traceability link. '
    'No orphaned elements detected.'
)

add_placeholder(
    'Insert: Cross-Model Traceability Matrix\n'
    '(MagicDraw: Matrix showing PIM elements on rows, CIM elements on columns,\n'
    'with relationship types color-coded. Recommended as a table diagram\n'
    'or exported matrix report from MagicDraw.)'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 15. TRANSFORMATION LOG
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('15. Transformation Log', level=1)

doc.add_heading('15.1 Verification Results', level=2)
add_table(
    ['Check', 'Status', 'Notes'],
    [
        ['CC-01: FR -> use case derivation', 'PASS', '40/41 linked; FR-040 justified orphan'],
        ['CC-02: Use case -> requirement coverage', 'WARNING', '14 non-normative scenarios not linked (by design)'],
        ['CC-04: PIM operation -> CIM use case', 'PASS', '15/15 traced'],
        ['CC-07: Action def parameters', 'PASS', 'All have in + out params'],
        ['CC-08: Interface def flows', 'PASS', '32 flows in 2 interfaces'],
        ['CC-10: CR regulatorySource attribute', 'ERROR', '10 CRs use doc text instead of formal attribute'],
        ['SC-09: No imports in nested packages', 'ERROR', 'CIM_Traceability has import statement'],
        ['NC-01..05: Naming conventions', 'PASS', 'PascalCase types, camelCase attrs, UPPER_SNAKE enums'],
        ['DC-01..07: Documentation', 'PASS', 'All elements have doc annotations in English'],
    ]
)

doc.add_heading('15.2 Assumptions', level=2)
assumptions = [
    'ASM-001: The IS SFM document is the sole authoritative source.',
    'ASM-002: Business scenarios (Section 3) are non-normative per ST-202.',
    'ASM-003: Security and authentication are delegated per ST-171.',
    'ASM-004: Exception message formats are delegated per ST-169.',
    'ASM-005: Data types are abstract at PIM level; ISO 21090 is a recommendation.',
    'ASM-006: Current version is the default per ST-177.',
    'ASM-007: Matching algorithms are configurable per organizational policy.',
    'ASM-008: Minimum state model (ACTIVE/INACTIVE) is mandatory.',
    'ASM-009: Automated merging is not encouraged per ST-128.',
]
for a in assumptions:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('15.3 Design Decisions', level=2)
decisions = [
    'DEC-001: CIM abstract attribute defs resolved to typed String attributes at PIM.',
    'DEC-002: PIM adds technical identifiers not present in CIM.',
    'DEC-003: PIM adds ServiceFault and FaultCategory for error handling.',
    'DEC-006: Two service interfaces (Management + Query) per SFM Section 2.3.1.',
    'DEC-007: Three consumer categories from SFM Section 3 actors.',
    'DEC-008: State machine includes "deprecated" beyond minimum ACTIVE/INACTIVE.',
    'DEC-009: Only 6 of 14 scenarios modeled as behavioral flows (multi-step only).',
]
for d in decisions:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('15.4 Open Issues', level=2)
add_table(
    ['Issue', 'Description', 'Owner'],
    [
        ['ISSUE-001', 'CC-10: 10 CR requirements lack formal regulatorySource attribute', 'SA3'],
        ['ISSUE-002', 'SC-09: CIM_Traceability contains import statement in nested package', 'SA3'],
    ]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX A: MODEL STATISTICS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Appendix A: Model Statistics', level=1)
add_table(
    ['Package', 'File', 'Lines', 'Element Types'],
    [
        ['CIM_IS.BusinessDomain', 'BusinessDomain.sysml', '582', '5 enum defs, 14 attribute defs, 13 item defs, 12 connection defs, 1 domain model'],
        ['CIM_IS.StakeholderModel', 'StakeholderModel.sysml', '133', '1 abstract part def, 9 part defs'],
        ['CIM_IS.BusinessCapabilities', 'BusinessCapabilities.sysml', '~1400', '29 use cases (15 normative + 14 scenarios)'],
        ['CIM_IS.BusinessRules', 'BusinessRules.sysml', '562', '42 constraint defs (6 formalized)'],
        ['CIM_IS.CIM_Requirements.FR', 'FunctionalRequirements.sysml', '373', '41 requirement usages'],
        ['CIM_IS.CIM_Requirements.QR', 'QualityRequirements.sysml', '159', '11 requirement usages'],
        ['CIM_IS.CIM_Requirements.CR', 'ComplianceRequirements.sysml', '147', '10 requirement usages'],
        ['CIM_IS.CIM_Traceability', 'CIM_Traceability.sysml', '635', '77 #derivation connections + diagnostics'],
        ['PIM_IS.DataModel', 'DataModel.sysml', '585', '45 item defs, 7 enum defs'],
        ['PIM_IS.ServiceContracts', 'ServiceContracts.sysml', '321', '2 port defs, 2 interface defs, 32 flows'],
        ['PIM_IS.Operations', 'Operations.sysml', '460', '15 action defs'],
        ['PIM_IS.BehavioralFlows', 'BehavioralFlows.sysml', '354', '6 action defs (workflow flows)'],
        ['PIM_IS.Composition', 'Composition.sysml', '262', '5 part defs, 1 state def, 6 interface connections'],
        ['PIM_IS.PIM_Traceability', 'PIM_Traceability.sysml', '619', '71 dependency statements'],
        ['Combined', 'IdentificationService_Complete.sysml', '6592', 'All of the above in single file'],
    ]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX B: GLOSSARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Appendix B: Glossary', level=1)
add_table(
    ['Term', 'Definition'],
    [
        ['CIM', 'Computation Independent Model - business-level abstraction without technical detail'],
        ['PIM', 'Platform Independent Model - design-level abstraction without platform specifics'],
        ['SFM', 'Service Functional Model - HL7 specification document type'],
        ['IS', 'Identification Service - the HL7 service being modeled'],
        ['PolicyDomain', 'An identity space or sphere of use (organizational, geographical, jurisdictional)'],
        ['EntityConcept', 'A semantic category (e.g., Person, Device) classifying EntityTypes'],
        ['EntityType', 'A specific entity type with schema and validation rules (semantic signifier)'],
        ['EntityTypeAssignment', 'Bridge customizing an EntityType for a specific PolicyDomain'],
        ['IdentityInstance', 'An actual identity entry with populated identifying information'],
        ['IdentityLink', 'A directed relationship between two IdentityInstances (link or merge)'],
        ['Semantic Signifier', 'The information structure (schema) defining identifying properties'],
        ['RWE', 'Real-World Entity - the actual thing being identified'],
        ['XIS', 'Cross-domain Identification Service (federated across organizations)'],
        ['MPI/EMPI', 'Master Patient Index / Enterprise Master Patient Index'],
        ['HIS', 'Hospital Information System - clinical system acting as service consumer'],
        ['RHIO', 'Regional Health Information Organization'],
    ]
)

# ═══════════════════════════════════════════════════════════════════
# APPENDIX C: DIAGRAM INDEX
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Appendix C: Diagram Placeholder Index', level=1)
doc.add_paragraph(
    'The following diagram placeholders are included in this document. After importing the SysML v2 '
    'model into MagicDraw 2026x, create these diagrams and replace the placeholders with the '
    'exported images.'
)
add_table(
    ['#', 'Diagram', 'Type', 'Section'],
    [
        ['1', 'Model Package Hierarchy', 'Package Diagram', '2'],
        ['2', 'Domain Entity Class Diagram', 'Block Definition Diagram', '3.3'],
        ['3', 'Domain Ontology Interconnection', 'Internal Block Diagram', '3.4'],
        ['4', 'Stakeholder Hierarchy', 'Block Definition Diagram', '4'],
        ['5', 'Normative Use Cases', 'Use Case Diagram', '5.2'],
        ['6', 'Business Scenarios', 'Use Case Diagram', '5.3'],
        ['7', 'Business Rules Dependencies', 'Requirement Diagram', '6'],
        ['8', 'Requirements Traceability Matrix', 'Requirements Table', '7.1'],
        ['9', 'CIM Traceability Matrix', 'Derivation Diagram', '8'],
        ['10', 'PIM Data Model Classes', 'Block Definition Diagram', '9'],
        ['11', 'Service Contract Interfaces', 'Interface Diagram', '10'],
        ['12', 'Operations Activity Diagram', 'Activity Diagram', '11'],
        ['13a-f', 'Behavioral Flow Diagrams (6)', 'Activity Diagrams', '12'],
        ['14', 'System Composition', 'Internal Block Diagram', '13.1'],
        ['15', 'Identity Lifecycle State Machine', 'State Machine Diagram', '13.2'],
        ['16', 'Cross-Model Traceability Matrix', 'Matrix / Table', '14'],
    ]
)

# ── Save ────────────────────────────────────────────────────────────
out_path = r'c:\Users\slotti\Documents\HL7-HSRA_E2-SFM2SysMLv2-Agent\output\ServiceFunctionalModel_IdentificationService\HL7_IS_SysMLv2_Model_Documentation.docx'
doc.save(out_path)
print(f'Document saved to: {out_path}')
